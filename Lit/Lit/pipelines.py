# Define your item pipelines here
#
# Don't forget to add your pipeline to the ITEM_PIPELINES setting
# See: https://docs.scrapy.org/en/latest/topics/item-pipeline.html


# useful for handling different item types with a single interface
from itemadapter import ItemAdapter
import re
import os
import io
from urllib.request import urlopen
import lxml.html
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .config import app_config

class AuthorPipeline:
    def process_item(self, item, spider):
        links_folder = app_config['LINKS_SAVE_PATH']
        author = item['author']
        links = "\n".join(item['links'])
        os.makedirs(links_folder, exist_ok=True)
        save_path = os.path.join(links_folder, author+".txt")
        with open(save_path,'w') as f:
            f.write(links)

        # print(item['author'])
        # print(item['links'])
        # print(item['dates'])
        # print(item['tags'])
        # print(item['title'])
        # print(item['content'])
        # print("\n")
        # return item
        

class DebugPipeline:

    def process_item(self, item, spider):
        
        print(item['tags'])
        print(item['title'])
        print(item['content'])
        print("\n")
        return item
 
        

class WordDocPipeline: 
    
    def __init__(self):
        self.save_path = app_config["STORIES_SAVE_PATH"]
        self.include_author_folder = app_config["INCLUDE_AUTHOR_FOLDER"]
        self.__item= None

    def close_spider(self, spider):
        print(self.__item['author'] + "  has been scraped")
        
    def process_item(self, item, spider):
        self.__item = item
        item['content'] = self.process_content(item)
        self.save(item) 
    
    def process_content(self, item):
        body = []
        for para in item['content']:
            if not para.startswith("<p><img"):
                clean_para = lxml.html.document_fromstring(para).text_content()
                body.append(clean_para)  
            else:
                body.append(para)
        return body

    
    def save(self, item):
        document = Document()
        document.add_heading(item['title'], 0)  

        for i in item['content']:
            if i.startswith('<p><img'):
                url = 'https://www.literotica.com'
                img = lxml.html.document_fromstring(i)
                img = img.xpath('//img/@src')[0]
                img = url + img
                document.add_picture(io.BytesIO(urlopen(img).read()))
                last_paragraph = document.paragraphs[-1] 
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else: 
                document.add_paragraph(i)
                paragraph_format = document.styles['Normal'].paragraph_format
                paragraph_format.space_after = Pt(12)
        
        if self.include_author_folder:
            save_path = os.path.join(self.save_path, item['author'])
        else:
            save_path = self.save_path
        os.makedirs(save_path, exist_ok = True)
        path = os.path.join(save_path, item['title']+".docx")
        document.save(path)

